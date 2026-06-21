"""HTML/Markdown report generation."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from openassay.errors import ReportError

DISCLAIMER = (
    "This report was generated using openassay (open-source). "
    "Final acceptance decisions and regulatory interpretation should be "
    "reviewed by qualified bioanalytical scientists."
)


def _missing_report_dependency(package: str) -> ReportError:
    return ReportError(
        f"{package} is required for this report format. "
        "Install optional report dependencies with: pip install 'openassay[reports]'."
    )


def generate_html_report(
    curve_result: Any,
    backcalc_results: list[Any],
    acceptance_result: Any,
    path: str,
) -> None:
    """Generate an HTML report.

    Parameters
    ----------
    curve_result : CalibrationResult
        Fitted standard curve result.
    backcalc_results : list
        List of BackCalcResult objects.
    acceptance_result : AcceptanceResult
        Acceptance evaluation result.
    path : str
        Output file path.
    """
    html_content = f"""<!DOCTYPE html>
<html>
<head>
    <title>openassay Run Report</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f2f2f2; }}
        .disclaimer {{ color: #666; font-style: italic; margin-top: 40px; }}
    </style>
</head>
<body>
    <h1>openassay Run Report</h1>
    <h2>Standard Curve</h2>
    <p>Model: {curve_result.fit_result.model_id}</p>
    <p>Weighting: {curve_result.fit_result.weight_scheme}</p>
    <p>R^2: {curve_result.fit_result.r_squared:.4f}</p>

    <h2>Parameters</h2>
    <table>
        <tr><th>Parameter</th><th>Value</th><th>SE</th></tr>
"""
    for name, val in curve_result.fit_result.params.items():
        se = curve_result.fit_result.se.get(name, float("nan"))
        html_content += f"        <tr><td>{name}</td><td>{val:.4g}</td><td>{se:.4g}</td></tr>\n"

    html_content += """    </table>

    <h2>Back-Calculated Samples</h2>
    <table>
        <tr>
            <th>Sample</th><th>Predicted</th><th>Diluted</th>
            <th>Below LLOQ</th><th>Above ULOQ</th>
        </tr>
"""
    for res in backcalc_results:
        html_content += (
            f"        <tr><td>{res.sample_name}</td><td>{res.predicted_concentration:.4g}</td>"
            f"<td>{res.diluted_concentration:.4g}</td><td>{res.below_lloq}</td><td>{res.above_uloq}</td></tr>\n"
        )

    html_content += f"""    </table>

    <h2>Acceptance</h2>
    <p>Passed: {acceptance_result.passed}</p>
    <ul>
"""
    for reason in acceptance_result.reasons:
        html_content += f"        <li>{reason}</li>\n"

    html_content += f"""    </ul>

    <p class="disclaimer">{DISCLAIMER}</p>
</body>
</html>
"""
    with open(path, "w", encoding="utf-8") as f:
        f.write(html_content)


def generate_markdown_report(
    curve_result: Any,
    backcalc_results: list[Any],
    acceptance_result: Any,
    path: str,
) -> None:
    """Generate a Markdown report.

    Parameters
    ----------
    curve_result : CalibrationResult
        Fitted standard curve result.
    backcalc_results : list
        List of BackCalcResult objects.
    acceptance_result : AcceptanceResult
        Acceptance evaluation result.
    path : str
        Output file path.
    """
    md_content = f"""# openassay Run Report

## Standard Curve
- Model: {curve_result.fit_result.model_id}
- Weighting: {curve_result.fit_result.weight_scheme}
- R^2: {curve_result.fit_result.r_squared:.4f}

## Parameters
| Parameter | Value | SE |
|-----------|-------|----|
"""
    for name, val in curve_result.fit_result.params.items():
        se = curve_result.fit_result.se.get(name, float("nan"))
        md_content += f"| {name} | {val:.4g} | {se:.4g} |\n"

    md_content += """
## Back-Calculated Samples
| Sample | Predicted | Diluted | Below LLOQ | Above ULOQ |
|--------|-----------|---------|------------|------------|
"""
    for res in backcalc_results:
        md_content += (
            f"| {res.sample_name} | {res.predicted_concentration:.4g} | "
            f"{res.diluted_concentration:.4g} | {res.below_lloq} | {res.above_uloq} |\n"
        )

    md_content += f"""
## Acceptance
- Passed: {acceptance_result.passed}
"""
    for reason in acceptance_result.reasons:
        md_content += f"  - {reason}\n"

    md_content += f"\n> *{DISCLAIMER}*\n"

    with open(path, "w", encoding="utf-8") as f:
        f.write(md_content)


def generate_pdf_report(
    curve_result: Any,
    backcalc_results: list[Any],
    acceptance_result: Any,
    path: str,
) -> None:
    """Generate a simple PDF report using optional ReportLab dependency."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError as exc:  # pragma: no cover - exercised with monkeypatch tests.
        raise _missing_report_dependency("reportlab") from exc

    pdf = canvas.Canvas(path, pagesize=letter)
    _, height = letter
    y = height - 72

    def line(text: str) -> None:
        nonlocal y
        if y < 72:
            pdf.showPage()
            y = height - 72
        pdf.drawString(72, y, text)
        y -= 16

    line("openassay Run Report")
    line(f"Model: {curve_result.fit_result.model_id}")
    line(f"Weighting: {curve_result.fit_result.weight_scheme}")
    line(f"R^2: {curve_result.fit_result.r_squared:.4f}")
    line("Parameters:")
    for name, val in curve_result.fit_result.params.items():
        se = curve_result.fit_result.se.get(name, float("nan"))
        line(f"- {name}: {val:.4g} (SE {se:.4g})")
    line("Back-Calculated Samples:")
    for res in backcalc_results:
        line(
            f"- {res.sample_name}: predicted {res.predicted_concentration:.4g}, "
            f"diluted {res.diluted_concentration:.4g}"
        )
    line(f"Acceptance passed: {acceptance_result.passed}")
    for reason in acceptance_result.reasons:
        line(f"- {reason}")
    line(DISCLAIMER)
    pdf.save()


def generate_docx_report(
    curve_result: Any,
    backcalc_results: list[Any],
    acceptance_result: Any,
    path: str,
) -> None:
    """Generate a simple DOCX report using optional python-docx dependency."""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - exercised with monkeypatch tests.
        raise _missing_report_dependency("python-docx") from exc

    document = Document()
    document.add_heading("openassay Run Report", level=1)
    document.add_heading("Standard Curve", level=2)
    document.add_paragraph(f"Model: {curve_result.fit_result.model_id}")
    document.add_paragraph(f"Weighting: {curve_result.fit_result.weight_scheme}")
    document.add_paragraph(f"R^2: {curve_result.fit_result.r_squared:.4f}")

    document.add_heading("Parameters", level=2)
    for name, val in curve_result.fit_result.params.items():
        se = curve_result.fit_result.se.get(name, float("nan"))
        document.add_paragraph(f"{name}: {val:.4g} (SE {se:.4g})")

    document.add_heading("Back-Calculated Samples", level=2)
    for res in backcalc_results:
        document.add_paragraph(
            f"{res.sample_name}: predicted {res.predicted_concentration:.4g}, "
            f"diluted {res.diluted_concentration:.4g}"
        )

    document.add_heading("Acceptance", level=2)
    document.add_paragraph(f"Passed: {acceptance_result.passed}")
    for reason in acceptance_result.reasons:
        document.add_paragraph(reason, style="List Bullet")
    document.add_paragraph(DISCLAIMER)
    document.save(path)


def report_run(
    curve_result: Any,
    backcalc_results: list[Any],
    acceptance_result: Any,
    path: str | Path,
    *,
    format: str = "auto",
) -> Path:
    """Generate a run report and return the output path."""
    output_path = Path(path)
    report_format = output_path.suffix.lower().lstrip(".") if format == "auto" else format.lower()

    if report_format in {"html", "htm"}:
        generate_html_report(curve_result, backcalc_results, acceptance_result, str(output_path))
    elif report_format in {"md", "markdown"}:
        generate_markdown_report(
            curve_result, backcalc_results, acceptance_result, str(output_path)
        )
    elif report_format == "pdf":
        generate_pdf_report(curve_result, backcalc_results, acceptance_result, str(output_path))
    elif report_format == "docx":
        generate_docx_report(curve_result, backcalc_results, acceptance_result, str(output_path))
    else:
        raise ValueError("format must be 'auto', 'html', 'markdown', 'pdf', or 'docx'")

    return output_path
